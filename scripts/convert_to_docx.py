"""
convert_to_docx.py — Convierte tesis_completa.md a DOCX usando plantilla_individual.docx
Ejecutar: python scripts/convert_to_docx.py
Requiere: python-docx (pip install python-docx)
"""

import os, re, pathlib, subprocess, sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = pathlib.Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "capitulos" / "tesis_completa.md"
TEMPLATE_PATH = ROOT / "plantilla_individual.docx"
OUTPUT_PATH = ROOT / "TFM_Julio_Rubio.docx"

if not MD_PATH.exists():
    sys.exit(f"No se encuentra {MD_PATH}")
if not TEMPLATE_PATH.exists():
    sys.exit(f"No se encuentra {TEMPLATE_PATH}")

with open(MD_PATH, "rb") as f:
    raw = f.read()
if raw[:3] == b"\xef\xbb\xbf":
    raw = raw[3:]
text_md = raw.decode("utf-8")

# Load template
doc = Document(str(TEMPLATE_PATH))

# ─── Limpiar la plantilla (borrar contenido placeholder) ───
# Keep only the first section (cover page)
while len(doc.paragraphs) > 0:
    p = doc.paragraphs[0]
    p._element.getparent().remove(p._element)
    if len(doc.paragraphs) == 0:
        break

# ─── Helper functions ───
def add_paragraph(text, style="Normal", bold=False, italic=False, font_size=None, alignment=None, font_name=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if alignment is not None:
        p.alignment = alignment
    return p


def add_heading(text, level=1):
    """Add a heading with the appropriate style."""
    if level == 1:
        style_name = "Heading 1"
    elif level == 2:
        style_name = "Heading 2"
    elif level == 3:
        style_name = "Heading 3"
    else:
        style_name = "Heading 4"
    
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text.strip())
    run.font.name = "Times New Roman"
    # Ensure heading numbering matches
    return p


def add_table_row(table, cells, bold=False):
    """Add a row to a table."""
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        if i < len(row.cells):
            row.cells[i].text = cell_text.strip()
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    if bold:
                        run.bold = True
    return row


# ─── Process markdown line by line ───
lines = text_md.split("\n")
i = 0
# Skip YAML front matter (--- ... ---)
if lines[0].strip() == "---":
    for j in range(1, len(lines)):
        if lines[j].strip() == "---":
            i = j + 1
            break

# Skip the YAML metadata lines
in_yaml = False
in_code = False
in_table = False
current_table = None
skip_yaml = True

while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    next_line = lines[i + 1] if i + 1 < len(lines) else ""
    next_stripped = next_line.strip()
    
    # Skip YAML front matter
    if stripped == "---" and skip_yaml:
        skip_yaml = False
        i += 1
        continue
    if skip_yaml:
        i += 1
        continue
    
    # Empty lines
    if stripped == "":
        i += 1
        continue
    
    # Horizontal rule (section break)
    if stripped == "---" and not skip_yaml:
        p = doc.add_paragraph()
        run = p.add_run("_" * 60)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        i += 1
        continue
    
    # Image references: ![Figura X.Y - Descripción](path)
    img_match = re.match(r"!\[(.*?)\]\((file:///)?([^)]+\.png)\)", line)
    if img_match:
        caption = img_match.group(1)
        img_path = img_match.group(3)
        # If it's a file:// URI, extract the path
        if img_path.startswith("file:///"):
            img_path = img_path[8:]
        # If it's a relative path, resolve it
        if not os.path.isabs(img_path):
            img_path = os.path.normpath(os.path.join(ROOT, "capitulos", img_path))
        
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Cm(14))
        else:
            # Placeholder
            add_paragraph(f"[Imagen no encontrada: {img_path}]")
        
        # Caption (next line usually has the *Figura...* caption)
        # If the next line is a figure caption, we'll handle it there
        i += 1
        continue
    
    # Figure captions: *Figura X.Y. Descripción...*
    if stripped.startswith("*Figura") or stripped.startswith("*Fuente"):
        p = doc.add_paragraph(style="Figuras" if "Figura" in stripped else "Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stripped.strip("*"))
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        run.italic = True
        i += 1
        continue
    
    # Headings
    if stripped.startswith("# ") and stripped[2:].strip():
        # Title or main heading: # Resumen, # Abstract, # CAPÍTULO 1
        text = stripped[2:].strip()
        if "CAPÍTULO" in text.upper() or "CAPITULO" in text.upper():
            add_heading(text, level=1)
        elif text.lower() in ("resumen", "abstract"):
            p = doc.add_paragraph(style="Título Índices" if "Título Índices" in [s.name for s in doc.styles] else "Normal")
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
        else:
            add_heading(text, level=1)
        i += 1
        continue
    
    if stripped.startswith("## ") and stripped[3:].strip():
        add_heading(stripped[3:].strip(), level=2)
        i += 1
        continue
    
    if stripped.startswith("### ") and stripped[4:].strip():
        add_heading(stripped[4:].strip(), level=3)
        i += 1
        continue
    
    if stripped.startswith("#### ") and stripped[5:].strip():
        add_heading(stripped[5:].strip(), level=4)
        i += 1
        continue
    
    # Bold text lines (like **Palabras clave:**, **Keywords:**)
    if stripped.startswith("**") and ":**" in stripped:
        p = doc.add_paragraph(style="Normal")
        parts = stripped.split(":**", 1)
        run_b = p.add_run(parts[0] + ":**")
        run_b.bold = True
        run_b.font.name = "Times New Roman"
        if len(parts) > 1:
            run_n = p.add_run(parts[1])
            run_n.font.name = "Times New Roman"
        i += 1
        continue
    
    # Tables
    if stripped.startswith("|") and stripped.endswith("|"):
        cells = [c.strip() for c in stripped.split("|")[1:-1]]
        
        # Skip separator lines (| --- | :--- |)
        if all(c in ("---", ":---", ":---:") for c in cells):
            i += 1
            continue
        
        if not in_table:
            # Start new table
            n_cols = len(cells)
            current_table = doc.add_table(rows=0, cols=n_cols)
            current_table.style = "Table Grid"
            in_table = True
        
        row = current_table.add_row()
        for j, cell_text in enumerate(cells):
            if j < len(row.cells):
                row.cells[j].text = cell_text
                for para in row.cells[j].paragraphs:
                    para.style = doc.styles["Normal"]
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Times New Roman"
                        # Detect header row (first row)
                        if len(current_table.rows) == 1:
                            run.bold = True
        i += 1
        continue
    else:
        if in_table:
            in_table = False
            current_table = None
    
    # Regular paragraphs
    # Process inline formatting: **bold** and *italic*
    p = doc.add_paragraph(style="Normal")
    
    # Parse inline formatting
    remaining = stripped
    while remaining:
        # Bold: **text**
        bold_match = re.search(r"\*\*(.+?)\*\*", remaining)
        # Italic: *text* (but not **)
        italic_match = re.search(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", remaining)
        
        if bold_match and (not italic_match or bold_match.start() < italic_match.start()):
            # Add text before bold
            if bold_match.start() > 0:
                run = p.add_run(remaining[:bold_match.start()])
                run.font.name = "Times New Roman"
            # Add bold text
            run = p.add_run(bold_match.group(1))
            run.bold = True
            run.font.name = "Times New Roman"
            remaining = remaining[bold_match.end():]
        elif italic_match:
            if italic_match.start() > 0:
                run = p.add_run(remaining[:italic_match.start()])
                run.font.name = "Times New Roman"
            run = p.add_run(italic_match.group(1))
            run.italic = True
            run.font.name = "Times New Roman"
            remaining = remaining[italic_match.end():]
        else:
            # No more formatting
            run = p.add_run(remaining)
            run.font.name = "Times New Roman"
            remaining = ""
    
    i += 1

# ─── Save ───
doc.save(str(OUTPUT_PATH))
print(f"Documento guardado: {OUTPUT_PATH}")
print(f"Tamaño: {OUTPUT_PATH.stat().st_size / 1024:.0f} KB")
